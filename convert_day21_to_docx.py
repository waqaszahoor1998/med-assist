#!/usr/bin/env python3
"""
Convert Day 21 Final Summary Report to DOCX format
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_day21_to_docx(md_file, docx_file):
    """Convert Day 21 markdown file to DOCX format"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set up styles
    setup_styles(doc)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('#').strip()
            
            if level == 1:
                add_header1(doc, header_text)
            elif level == 2:
                add_header2(doc, header_text)
            elif level == 3:
                add_header3(doc, header_text)
            elif level == 4:
                add_header4(doc, header_text)
                
        # Handle bold text
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            add_bold_paragraph(doc, text)
            
        # Handle bullet points
        elif line.startswith('- '):
            text = line[2:].strip()
            add_bullet_point(doc, text)
            
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            add_numbered_point(doc, text)
            
        # Handle code blocks
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
            
        # Handle regular paragraphs
        else:
            if line:
                add_paragraph(doc, line)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def setup_styles(doc):
    """Set up document styles"""
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Inches(0.25)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header 1 style
    h1_style = doc.styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Inches(0.2)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Inches(0.1)
    h1_style.paragraph_format.space_after = Inches(0.05)
    
    # Header 2 style
    h2_style = doc.styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Inches(0.18)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Inches(0.08)
    h2_style.paragraph_format.space_after = Inches(0.04)
    
    # Header 3 style
    h3_style = doc.styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = 'Arial'
    h3_style.font.size = Inches(0.16)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Inches(0.06)
    h3_style.paragraph_format.space_after = Inches(0.03)
    
    # Header 4 style
    h4_style = doc.styles.add_style('CustomH4', WD_STYLE_TYPE.PARAGRAPH)
    h4_style.font.name = 'Arial'
    h4_style.font.size = Inches(0.14)
    h4_style.font.bold = True
    h4_style.paragraph_format.space_before = Inches(0.04)
    h4_style.paragraph_format.space_after = Inches(0.02)

def add_header1(doc, text):
    """Add header 1"""
    p = doc.add_paragraph(text, style='CustomH1')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_header2(doc, text):
    """Add header 2"""
    p = doc.add_paragraph(text, style='CustomH2')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_header3(doc, text):
    """Add header 3"""
    p = doc.add_paragraph(text, style='CustomH3')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_header4(doc, text):
    """Add header 4"""
    p = doc.add_paragraph(text, style='CustomH4')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_bold_paragraph(doc, text):
    """Add bold paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Inches(0.12)

def add_bullet_point(doc, text):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.style.font.name = 'Arial'
    p.style.font.size = Inches(0.11)

def add_numbered_point(doc, text):
    """Add numbered point"""
    p = doc.add_paragraph(text, style='List Number')
    p.style.font.name = 'Arial'
    p.style.font.size = Inches(0.11)

def add_code_block(doc, code):
    """Add code block"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Inches(0.1)
    p.paragraph_format.left_indent = Inches(0.5)

def add_paragraph(doc, text):
    """Add regular paragraph"""
    p = doc.add_paragraph(text)
    p.style.font.name = 'Arial'
    p.style.font.size = Inches(0.11)

if __name__ == "__main__":
    # Convert the Day 21 report to DOCX
    convert_day21_to_docx(
        '/Users/m.w.zahoor/Desktop/med assist/docs/day21-final-summary-report.md',
        '/Users/m.w.zahoor/Desktop/med assist/docs/day21-final-summary-report.docx'
    )
